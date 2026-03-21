from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

def generate_report(project_data):
    doc = Document()
    
    # Title Page
    doc.add_heading(project_data.get('title', 'Project Report'), 0)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"\nA Project Report Submitted for {project_data.get('domain', 'Computer Science')}\n").bold = True
    
    doc.add_page_break()
    
    # Abstract
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph(project_data.get('abstract', 'No abstract provided.'))
    
    # Problem Statement
    doc.add_heading('1. Problem Statement', level=1)
    doc.add_paragraph(project_data.get('problem_statement', ''))

    # Literature Survey
    if project_data.get('literature_survey'):
        doc.add_heading('2. Literature Survey', level=1)
        doc.add_paragraph(project_data.get('literature_survey', ''))
    
    # Features
    if project_data.get('features'):
        doc.add_heading('3. Key Features', level=1)
        for feature in project_data.get('features', []):
            doc.add_paragraph(feature, style='List Bullet')

    # Architecture
    doc.add_heading('4. System Architecture', level=1)
    doc.add_paragraph(project_data.get('architecture_description', ''))
    
    if project_data.get('database_design'):
        doc.add_heading('5. Database Design', level=1)
        doc.add_paragraph(project_data.get('database_design', ''))

    if project_data.get('logic_flow'):
        doc.add_heading('6. Logic Flow & Security', level=1)
        doc.add_paragraph(project_data.get('logic_flow', ''))
        if project_data.get('security_measures'):
            doc.add_paragraph(project_data.get('security_measures', ''))

    # Technical Stack
    doc.add_heading('7. Technology Stack', level=1)
    tech = project_data.get('tech_stack_details', {})
    for k, v in tech.items():
        doc.add_paragraph(f"{k}: {v}", style='List Bullet')
    
    # Methodology
    if project_data.get('methodology'):
        doc.add_heading('8. Methodology', level=1)
        doc.add_paragraph(project_data.get('methodology', ''))

    # Conclusion
    doc.add_heading('9. Conclusion', level=1)
    doc.add_paragraph("This project demonstrates a practical implementation of the proposed system with professional architecture and robust implementation highlights.")
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
